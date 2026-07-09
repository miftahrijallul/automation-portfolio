import docx

from jats_converter.docx_reader import read_docx


def make_docx(path):
    d = docx.Document()
    d.add_heading("Introduction", level=1)
    d.add_paragraph("First paragraph of the introduction.")
    d.add_paragraph("Second paragraph.")
    d.add_heading("References", level=1)
    d.add_paragraph("Kari, Z. A. et al. (2022). Phytobiotics. Journal X.")
    d.add_paragraph("Celano, R. et al. (2021). Onion flavonoids. Journal Y.")
    d.save(str(path))


def make_docx_bold_headings(path):
    """Naskah gaya umum Indonesia: heading = paragraf Normal yang di-bold."""
    d = docx.Document()
    d.add_paragraph("Effect of Something Long Enough to Not Be a Heading Because It Runs On").add_run()
    p = d.add_paragraph()
    p.add_run("INTRODUCTION").bold = True
    d.add_paragraph("Intro paragraph one.")
    p = d.add_paragraph()
    p.add_run("Keywords: additive, broiler, onion peel, production performance").bold = True
    p = d.add_paragraph()
    p.add_run("MATERIALS AND METHODS").bold = True
    d.add_paragraph("Methods paragraph.")
    p = d.add_paragraph()
    p.add_run("REFERENCES").bold = True
    d.add_paragraph("Someone, A. (2020). A cited work. Journal Z.")
    d.save(str(path))


def test_read_docx_detects_bold_paragraph_headings(tmp_path):
    f = tmp_path / "bold.docx"
    make_docx_bold_headings(f)
    sections, refs = read_docx(str(f))

    titles = [t for t, _ in sections]
    assert titles == ["INTRODUCTION", "MATERIALS AND METHODS"]
    assert sections[0][1] == [
        "Intro paragraph one.",
        "Keywords: additive, broiler, onion peel, production performance",
    ]
    assert refs == ["Someone, A. (2020). A cited work. Journal Z."]


def test_read_docx_splits_sections_and_refs(tmp_path):
    f = tmp_path / "manuscript.docx"
    make_docx(f)
    sections, refs = read_docx(str(f))

    assert len(sections) == 1
    title, paragraphs = sections[0]
    assert title == "Introduction"
    assert paragraphs == [
        "First paragraph of the introduction.",
        "Second paragraph.",
    ]
    assert len(refs) == 2
    assert refs[0].startswith("Kari, Z. A.")
